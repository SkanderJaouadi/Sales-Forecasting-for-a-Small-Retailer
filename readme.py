try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: Please run 'pip install python-docx' first.")
    exit()

# Initialize Document
doc = Document()

# Title
title = doc.add_heading('🛒 Grocery Sales Forecasting System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Quick Start Section
doc.add_heading('🚀 Getting Started', level=1)
p = doc.add_paragraph()
p.add_run('1. Install Dependencies:').bold = True
doc.add_paragraph('pip install pandas numpy xgboost scikit-learn', style='No Spacing')

p = doc.add_paragraph()
p.add_run('2. Prepare Your Data:').bold = True
doc.add_paragraph('Ensure grocery_chain_data.csv is in the root folder.', style='No Spacing')

p = doc.add_paragraph()
p.add_run('3. Train the Model:').bold = True
doc.add_paragraph('python train_model.py', style='No Spacing')

p = doc.add_paragraph()
p.add_run('4. Predict Sales:').bold = True
doc.add_paragraph('python predict_sales.py', style='No Spacing')

# How It Works Section
doc.add_heading('🛠️ How It Works', level=1)
doc.add_paragraph(
    "The system processes historical transaction data to create time-based features "
    "(lags, rolling means, and calendar effects) to forecast demand. The interactive "
    "script is case-insensitive, meaning you can type 'apples' or 'Apples' and get "
    "the same accurate result."
)

# Model Details Section
doc.add_heading('📊 Model Details', level=1)
items = [
    'Algorithm: XGBoost Regressor',
    'Output: Predicted quantity (rounded to 2 decimal places)',
    'Features: Product/Store encoding, seasonality, and historical sales trends.'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# Save the document
file_name = "README.docx"
doc.save(file_name)

print(f"✅ {file_name} has been created successfully!")